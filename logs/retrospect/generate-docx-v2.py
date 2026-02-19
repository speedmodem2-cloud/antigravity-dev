import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# PAGE SETUP (match reference: Letter, margins)
# ============================================================
section = doc.sections[0]
section.page_width = 7772400   # EMU
section.page_height = 10058400
section.top_margin = 914400
section.bottom_margin = 914400
section.left_margin = 762000
section.right_margin = 762000

# ============================================================
# STYLES (match reference exactly)
# ============================================================
# Heading 1: Arial, 16pt, #1E3A5F, bold
h1 = doc.styles['Heading 1']
h1.font.name = 'Arial Unicode MS'
h1.font.size = Pt(16)
h1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
h1.font.bold = True
h1.paragraph_format.space_before = Pt(10)
h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.line_spacing = 1.0

# Heading 2: Arial, 12pt, #2563EB, bold
h2 = doc.styles['Heading 2']
h2.font.name = 'Arial Unicode MS'
h2.font.size = Pt(12)
h2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
h2.font.bold = True
h2.paragraph_format.space_before = Pt(8)
h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.line_spacing = 1.0

# Normal
normal = doc.styles['Normal']
normal.font.name = 'Arial Unicode MS'
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(3)

# ============================================================
# HELPERS
# ============================================================
def add_run(para, text, size=10, bold=False, color=None, font='Arial Unicode MS'):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(table):
    """Set table borders matching reference (thin black)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'CCCCCC')
        borders.append(element)
    tbl_pr.append(borders)

def add_styled_table(headers, rows, col_widths=None):
    """Create table matching reference style: #1E3A5F header, white text."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '1e3a5f')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Arial Unicode MS'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Alternate row shading
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'f8fafc')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Arial Unicode MS'
            run.font.size = Pt(10)
            # First column bold
            if c_idx == 0:
                run.bold = True

    return table

def add_kv_table(kv_pairs):
    """2-column key-value table matching reference."""
    table = doc.add_table(rows=len(kv_pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    for i, (key, val) in enumerate(kv_pairs):
        # Key cell
        cell_k = table.rows[i].cells[0]
        set_cell_shading(cell_k, 'f1f5f9')
        p = cell_k.paragraphs[0]
        run = p.add_run(key)
        run.font.name = 'Arial Unicode MS'
        run.font.size = Pt(10)
        run.bold = True

        # Value cell
        cell_v = table.rows[i].cells[1]
        if i % 2 == 0:
            set_cell_shading(cell_v, 'ffffff')
        else:
            set_cell_shading(cell_v, 'f8fafc')
        p = cell_v.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Arial Unicode MS'
        run.font.size = Pt(10)

    return table

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True)
        add_run(p, text)
    else:
        add_run(p, f'• {text}')

# ============================================================
# COVER PAGE (match reference layout)
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

# Company / Context
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'AntiGravity Dev Environment', size=14, color=(0x64, 0x74, 0x8b))

# Main Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'P1~P4 프로젝트', size=24, bold=True, color=(0x1E, 0x3A, 0x5F))

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '통합 회고 보고서', size=20, bold=True, color=(0x25, 0x63, 0xEB))

# Description
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '프론트엔드 3개 + 백엔드 1개 · 토큰 비용 · 오케스트레이션 · KPT', size=11, color=(0x64, 0x74, 0x8b))

doc.add_paragraph('')

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '2026년 2월 18일', size=11, color=(0x94, 0xA3, 0xB8))

# Project
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Opus 오케스트레이션 + Sonnet 서브에이전트 체제', size=10, color=(0x94, 0xA3, 0xB8))

doc.add_page_break()

# ============================================================
# TOC (match reference: numbered, blue, bold)
# ============================================================
doc.add_heading('목차', level=1)
doc.add_paragraph('')

toc_items = [
    '프로젝트 총괄',
    '토큰 비용 분석',
    '기술 성장 궤적',
    '오케스트레이션 진화',
    '반복된 문제 (4건)',
    'KPT 종합',
    '성과 요약',
]
for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    add_run(p, f'{i}. ', size=11, bold=True, color=(0x25, 0x63, 0xEB))
    add_run(p, item, size=11)

doc.add_page_break()

# ============================================================
# 1. 프로젝트 총괄
# ============================================================
doc.add_heading('1. 프로젝트 총괄', level=1)
p = doc.add_paragraph()
add_run(p, '2일간(02-17~02-18) 4개 프로젝트를 완료했습니다. 프론트엔드 포트폴리오 3종과 NestJS 백엔드 1종입니다.')

add_styled_table(
    ['#', '프로젝트', '유형', '디자인/스택', '날짜', 'Phase', 'dist'],
    [
        ['P1', 'portfolio', '프론트엔드', 'Synthwave', '02-17', '7/7', '17MB*'],
        ['P2', 'portfolio2', '프론트엔드', 'Brutalist', '02-18', '6/7', '548KB'],
        ['P3', 'portfolio3', '프론트엔드', 'Synthwave Enhanced', '02-18', '8/8', '2.2MB'],
        ['P4', 'shop-backend1', '백엔드', 'NestJS+MySQL+Docker', '02-18', '7/7', '863KB'],
    ]
)
doc.add_paragraph('')
add_bullet('P1 dist에 source map(.js.map) + PNG 이미지 포함. 실제 배포 시 ~3MB 이하')

# ============================================================
# 2. 토큰 비용 분석
# ============================================================
doc.add_heading('2. 토큰 비용 분석', level=1)
p = doc.add_paragraph()
add_run(p, '전체 4개 프로젝트의 AI 모델별 토큰 사용량과 비용입니다.')

add_styled_table(
    ['프로젝트', 'Opus', 'Sonnet', 'Gemini', 'Haiku', '합계'],
    [
        ['P1', '-', '-', '-', '-', '~$4 (추정)'],
        ['P2', '$1.64', '$2.07', '$0.13', '$0.03', '$3.87'],
        ['P3', '$1.64', '$2.07', '$0.13', '$0.03', '$3.87'],
        ['P4', '$4.44', '$1.78', '-', '-', '$6.22'],
        ['합계', '~$7.72', '~$5.92', '$0.26', '$0.06', '~$17.96'],
    ]
)

doc.add_heading('모델별 비중', level=2)
add_bullet('Opus: ~43% (P4 단독이면 71%)')
add_bullet('Sonnet: ~33%')
add_bullet('Gemini + Haiku: ~2%')

doc.add_heading('프로젝트별 효율', level=2)
add_bullet('P2가 가장 효율적: $3.87로 6컴포넌트, CSS 2레이어, 548KB dist')
add_bullet('P4가 가장 비효율적: Opus 지휘 비용이 전체의 71% ($4.44/$6.22)')

# ============================================================
# 3. 기술 성장 궤적
# ============================================================
doc.add_heading('3. 기술 성장 궤적', level=1)
p = doc.add_paragraph()
add_run(p, '각 프로젝트에서 새로 배운 핵심 기술 목록입니다.')

add_styled_table(
    ['프로젝트', '핵심 기술'],
    [
        ['P1', 'Vite, React, CSS 애니메이션, WebP 이미지 최적화'],
        ['P2', 'CSS-only 비주얼 레이어 (ghost text, dot patterns), 타이포그래피 시스템'],
        ['P3', 'CSS blend-mode, fixed bg overlay, ambient glow orbs, 이미지 전략 결정'],
        ['P4', 'NestJS 모듈 시스템, TypeORM 관계, JWT 비동기 등록, 트랜잭션, E2E 테스트'],
    ]
)

# ============================================================
# 4. 오케스트레이션 진화
# ============================================================
doc.add_heading('4. 오케스트레이션 진화', level=1)
p = doc.add_paragraph()
add_run(p, 'AI 에이전트 오케스트레이션 모델이 프로젝트를 거치며 어떻게 변화했는지 분석합니다.')

doc.add_heading('4-1. P1~P3: 프론트엔드 파이프라인', level=2)
add_bullet('Opus 지휘 → Sonnet 구현 → Gemini 리뷰 → Haiku 문서')
add_bullet('4모델 협업, 균형 잡힌 비용 분배')
add_bullet('Phase 스킵 문제 발견 → P2에서 확인 규칙 수립')

doc.add_heading('4-2. P4: 백엔드 전환', level=2)
add_bullet('Gemini/Haiku 미활용 → Opus+Sonnet 2모델 체제')
add_bullet('Opus 비중 급증 (71%) → 지휘 비용 절감 필요성 확인')
add_bullet('병렬 디스패치 패턴 확립 (Phase당 2~3 Sonnet 에이전트)')
add_bullet('서브에이전트 14세션 운용 (P2~P3 대비 2.8배)')

doc.add_heading('4-3. Opus 역할 축소 방향', level=2)
add_styled_table(
    ['구분', '역할', '비중'],
    [
        ['현재 (P4)', '판단 + 리뷰 + 디버깅 + 재디스패치', '71%'],
        ['목표 (P5~)', 'Phase 전환 + 최종 리뷰만', '<20%'],
    ]
)

# ============================================================
# 5. 반복된 문제
# ============================================================
doc.add_heading('5. 반복된 문제', level=1)
p = doc.add_paragraph()
add_run(p, '4개 프로젝트에서 반복적으로 발생한 문제와 해결 방안입니다.')

doc.add_heading('5-1. 대형 커밋 (전 프로젝트)', level=2)
add_kv_table([
    ('현상', '모든 프로젝트에서 10+ 파일 커밋 발생'),
    ('원인', '서브에이전트가 모듈 전체를 한 번에 생성'),
    ('해결', '"모듈당 1커밋" 규칙을 서브에이전트 프롬프트에 포함'),
    ('영향도', '★★★★☆'),
])

doc.add_heading('5-2. Phase 스킵 (P2)', level=2)
add_kv_table([
    ('현상', 'Phase 4 리뷰를 확인 없이 건너뜀'),
    ('원인', 'Phase skip 확인 절차 미수립'),
    ('해결', '"Phase skip = always confirm" 규칙 수립 → P3~P4에서 준수'),
    ('영향도', '★★★☆☆'),
])

doc.add_heading('5-3. 하드코딩 값 (P4)', level=2)
add_kv_table([
    ('현상', '테스트 이메일 고정 → DB 중복 실패. JWT 시크릿 타이밍 미스매치'),
    ('원인', '환경 의존 값을 리터럴로 사용'),
    ('해결', '동적 값 생성 (Date.now()) + JwtModule.registerAsync() 패턴'),
    ('영향도', '★★★★☆'),
])

doc.add_heading('5-4. dist 크기 관리 (P1)', level=2)
add_kv_table([
    ('현상', 'P1에서 17MB (source map + 이미지)'),
    ('원인', '빌드 시 .map 파일 포함, PNG 원본 사용'),
    ('해결', '.map 제외, WebP 이미지, no-image 전략 → P2~P3에서 548KB~2.2MB로 개선'),
    ('영향도', '★★★☆☆'),
])

# ============================================================
# 6. KPT 종합
# ============================================================
doc.add_heading('6. KPT 종합', level=1)

doc.add_heading('Keep (유지)', level=2)
keeps = [
    'Conventional Commits — 전 프로젝트 준수',
    '병렬 서브에이전트 디스패치 — Phase당 2~3 Sonnet 동시 실행',
    'Phase 기반 파이프라인 — 체계적 진행 + TUI 모니터링',
    '회고 + MEMORY 기록 — 교훈이 다음 프로젝트에 반영됨',
    '디자인 다양화 — synthwave → brutalist → enhanced, 매번 새 도전',
]
for k in keeps:
    add_bullet(k)

doc.add_heading('Problem (문제)', level=2)
problems = [
    'Opus 비용 과다 — P4에서 71%, 전체 ~43%',
    '대형 커밋 — 4개 프로젝트 모두 동일 문제',
    'retrospect 도구 한계 — 프론트엔드 지표만 (CSS layer, component 수)',
    'P1 회고 누락 — 초기 프로젝트라 retrospect 미실행',
]
for pr in problems:
    add_bullet(pr)

doc.add_heading('Try (시도)', level=2)
tries = [
    'Opus <20% 목표 — P5(shop-backend2)에서 적용',
    '모듈당 1커밋 규칙 — 서브에이전트 프롬프트에 명시',
    'retrospect 도구 확장 — 백엔드 지표 (API 수, 엔티티 수, 테스트 커버리지)',
    'E2E 선행 작성 — 구현 전 테스트 스켈레톤 생성',
    '4모델 체제 복원 — 백엔드에서도 Gemini 리뷰 + Haiku 문서 활용',
]
for t in tries:
    add_bullet(t)

# ============================================================
# 7. 성과 요약
# ============================================================
doc.add_heading('7. 성과 요약', level=1)

add_styled_table(
    ['지표', '값'],
    [
        ['기간', '02-17 ~ 02-18'],
        ['완료 프로젝트', '4개'],
        ['총 Phase 수', '29개 (스킵 1 포함)'],
        ['프론트엔드 컴포넌트', '16+'],
        ['백엔드 API 엔드포인트', '~20개'],
        ['백엔드 엔티티', '5개'],
        ['E2E 테스트', '31개 (전 통과)'],
        ['총 토큰 비용', '~$18'],
        ['다음 프로젝트', 'shop-backend2 (실시간 경매, 계획 완료)'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\Dev\system\docs\통합회고보고서_P1-P4.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
